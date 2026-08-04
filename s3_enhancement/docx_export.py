"""The Word export for both S3 hand-off documents.

Reads the same `docshell.Document` the HTML and PDF are built from, so the
three outputs cannot say different things. What Word gets that the other two do
not is the machinery a controlled document is expected to carry: real heading
styles, so the navigation pane works and a TOC field can find them; a TOC field,
so the page numbers are Word's rather than a guess; and PAGE/NUMPAGES fields in
the footer.

Three things are worth knowing before changing this module.

**Headings must use the built-in styles, not just look like them.** A run
coloured maroon at 15pt is not a heading to Word — it will not appear in the
navigation pane and `TOC \\o "1-3"` will not collect it. So each part and section
gets `Heading 1` / `Heading 2` and the styles themselves are recoloured once, at
the top, rather than the runs being formatted individually.

**The TOC is a field, and fields are not evaluated by python-docx.** The
document sets `w:updateFields` so Word offers to update on open; until then the
placeholder text under it says so, which is exactly what the client's own
template does. Nothing here can compute a page number, and nothing here
pretends to.

**Figures are rasterised, and say so when they cannot be.** The change map is
inline SVG. Word 2016+ can display SVG but python-docx has no API for it, so the
SVG goes through the Chromium that already renders the PDF and comes back as a
PNG. Chromium is optional on a locked-down host (see `designdoc.render_pdf`), so
a missing browser drops a stated line into the document rather than silently
losing the diagram — a design document that quietly ships without its change map
is worse than one that says the change map is missing.
"""

from __future__ import annotations

import io
import re

from docx import Document as _WordDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

from s3_enhancement import docshell
from s3_enhancement.docshell import (
    Bullets,
    Callout,
    Document,
    Figure,
    Para,
    Sub,
    Table,
)

BODY_FONT = "Aptos"
BODY_FALLBACK = "Calibri"
MONO_FONT = "Consolas"

INK = RGBColor(0x1F, 0x1F, 0x1F)
INK_SOFT = RGBColor(0x4D, 0x4D, 0x4D)
INK_FAINT = RGBColor(0x7A, 0x7A, 0x7A)
ACCENT = RGBColor(0x8B, 0x1E, 0x2D)
OK = RGBColor(0x1C, 0x6B, 0x3C)

TONE_COLOR = {"ok": OK, "bad": ACCENT, "todo": ACCENT}

# A4 portrait with the same margins the PDF uses, in EMU (914400 per inch).
_MM = 36000
PAGE_W, PAGE_H = 210 * _MM, 297 * _MM
MARGIN_X, MARGIN_TOP, MARGIN_BOTTOM = 18 * _MM, 20 * _MM, 20 * _MM
CONTENT_W = PAGE_W - 2 * MARGIN_X


# --- low-level OOXML helpers -------------------------------------------------


def _el(tag: str, **attrs) -> OxmlElement:
    node = OxmlElement(tag)
    for key, value in attrs.items():
        node.set(qn(f"w:{key}"), str(value))
    return node


def _field(paragraph, instruction: str, placeholder: str) -> None:
    """A Word field — `{ PAGE }`, `{ TOC ... }` — with text shown until updated."""
    begin = paragraph.add_run()._r
    begin.append(_el("w:fldChar", fldCharType="begin"))

    instr = paragraph.add_run()._r
    node = OxmlElement("w:instrText")
    node.set(qn("xml:space"), "preserve")
    node.text = instruction
    instr.append(node)

    separate = paragraph.add_run()._r
    separate.append(_el("w:fldChar", fldCharType="separate"))

    paragraph.add_run(placeholder)

    end = paragraph.add_run()._r
    end.append(_el("w:fldChar", fldCharType="end"))


def _shade(cell, fill: str) -> None:
    cell._tc.get_or_add_tcPr().append(_el("w:shd", val="clear", color="auto", fill=fill))


def _borders(table) -> None:
    """Hairline borders all round, matching the HTML's 0.75pt rules."""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el(f"w:{edge}", val="single", sz="4", space="0", color="D4D4D4"))
    table._tbl.tblPr.append(borders)


def _left_rule(table) -> None:
    """The cover's vertical maroon rule, as a table's left edge.

    Word draws no free-standing vertical lines, so the rule beside the
    identity block is a border on the block that sits against it — which is
    also how the HTML does it.
    """
    borders = OxmlElement("w:tblBorders")
    borders.append(_el("w:left", val="single", sz="18", space="0", color=docshell.ACCENT))
    for edge in ("top", "bottom", "right", "insideH", "insideV"):
        borders.append(_el(f"w:{edge}", val="none", sz="0", space="0", color="auto"))
    table._tbl.tblPr.append(borders)


def _rule_below(paragraph, color: str = "E6E6E6") -> None:
    borders = OxmlElement("w:pBdr")
    borders.append(_el("w:bottom", val="single", sz="4", space="4", color=color))
    paragraph._p.get_or_add_pPr().append(borders)


def _rule_above(paragraph, color: str = "E6E6E6") -> None:
    borders = OxmlElement("w:pBdr")
    borders.append(_el("w:top", val="single", sz="4", space="4", color=color))
    paragraph._p.get_or_add_pPr().append(borders)


def _left_border(paragraph) -> None:
    borders = OxmlElement("w:pBdr")
    borders.append(_el("w:left", val="single", sz="18", space="10", color=docshell.ACCENT))
    paragraph._p.get_or_add_pPr().append(borders)


def _clear(paragraph):
    """Empty a paragraph without leaving a run behind.

    `paragraph.text = ""` does not clear it — python-docx replaces the runs
    with a single empty one, which then sits at `runs[0]` carrying none of the
    formatting the real first run has. That is invisible until something reads
    the colour off `runs[0]` and finds nothing.
    """
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    return paragraph


def _keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _update_fields_on_open(document) -> None:
    settings = document.settings.element
    settings.append(_el("w:updateFields", val="true"))


# --- styles ------------------------------------------------------------------


def _configure_styles(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    # East-Asian and fallback faces, which the python-docx API does not reach.
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    rpr.set(qn("w:cs"), BODY_FALLBACK)
    rpr.set(qn("w:eastAsia"), BODY_FALLBACK)

    for name, size, color, before in (
        ("Heading 1", 16, ACCENT, 18),
        ("Heading 2", 13, ACCENT, 14),
        ("Heading 3", 11, INK, 10),
    ):
        style = document.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def _configure_page(document) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Emu(PAGE_W), Emu(PAGE_H)
    section.left_margin = section.right_margin = Emu(MARGIN_X)
    section.top_margin, section.bottom_margin = Emu(MARGIN_TOP), Emu(MARGIN_BOTTOM)
    section.header_distance = section.footer_distance = Emu(11 * _MM)
    return section


def _running_chrome(section, running_title: str) -> None:
    head = _clear(section.header.paragraphs[0])
    run = head.add_run(running_title)
    run.font.size = Pt(8)
    run.font.color.rgb = INK_FAINT
    _rule_below(head)

    foot = _clear(section.footer.paragraphs[0])
    _rule_above(foot)
    # A dot leader between the classification and the page number, which is
    # what the client's own footer does.
    foot.paragraph_format.tab_stops.add_tab_stop(
        Emu(CONTENT_W), WD_ALIGN_PARAGRAPH.RIGHT, WD_TAB_LEADER.DOTS
    )
    left = foot.add_run(f"{docshell.ORG} · {docshell.SYSTEM} — {docshell.CLASSIFICATION}\t")
    left.font.size = Pt(8)
    left.font.color.rgb = INK_FAINT
    foot.add_run("Page ").font.size = Pt(8)
    _field(foot, " PAGE ", "1")
    foot.add_run(" of ").font.size = Pt(8)
    _field(foot, " NUMPAGES ", "1")
    for run in foot.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = INK_FAINT


# --- blocks ------------------------------------------------------------------


def _write_runs(paragraph, items, *, size: float = 10.5, color=INK, italic: bool = False):
    for item in items:
        # Word has no <br>; a newline inside a run is a real line break element.
        for index, line in enumerate(item.text.split("\n")):
            if index:
                paragraph.add_run()._r.append(OxmlElement("w:br"))
            run = paragraph.add_run(line)
            run.bold = item.bold
            run.italic = italic
            run.font.size = Pt(size - 0.5 if item.code else size)
            run.font.name = MONO_FONT if item.code else BODY_FONT
            run.font.color.rgb = TONE_COLOR.get(item.tone, color)


def _write_table(document, block: Table) -> None:
    if not block.rows:
        return
    table = document.add_table(rows=1, cols=len(block.headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _borders(table)

    widths = block.widths or tuple([round(100 / len(block.headers))] * len(block.headers))
    for cell, header, width in zip(table.rows[0].cells, block.headers, widths, strict=False):
        cell.width = Emu(int(CONTENT_W * width / 100))
        _shade(cell, docshell.ACCENT)
        paragraph = _clear(cell.paragraphs[0])
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    table.rows[0]._tr.get_or_add_trPr().append(_el("w:tblHeader", val="true"))

    for row in block.rows:
        cells = table.add_row().cells
        for cell, content, width in zip(cells, row, widths, strict=False):
            cell.width = Emu(int(CONTENT_W * width / 100))
            _write_runs(_clear(cell.paragraphs[0]), content, size=9.5)
    document.add_paragraph()


def _write_figure(document, block: Figure) -> None:
    png = _rasterise(block.svg)
    if png is None:
        paragraph = document.add_paragraph()
        _write_runs(
            paragraph,
            docshell.runs(
                "[Change map not embedded — this host has no browser to render the "
                "diagram with. It is present in the PDF and HTML exports.]"
            ),
            size=9.5,
            color=ACCENT,
            italic=True,
        )
        return
    document.add_picture(io.BytesIO(png), width=Emu(int(CONTENT_W * 0.82)))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if block.caption:
        caption = document.add_paragraph()
        _write_runs(caption, docshell.runs(block.caption), size=9, color=INK_FAINT, italic=True)


_VIEWBOX_RE = re.compile(r'viewBox="\s*([\d.\-]+)[ ,]+([\d.\-]+)[ ,]+([\d.]+)[ ,]+([\d.]+)')


def _rasterise(svg: str, *, width: int = 1400) -> bytes | None:
    """SVG to PNG through the Chromium the PDF path already uses.

    Returns None when the browser is unavailable, which is a supported state on
    a locked-down host — the caller writes a line saying the figure is missing.
    """
    match = _VIEWBOX_RE.search(svg)
    ratio = (float(match.group(4)) / float(match.group(3))) if match else 0.5
    height = max(1, round(width * ratio))
    try:
        from playwright.sync_api import Error as PlaywrightError
        from playwright.sync_api import sync_playwright
    except ImportError:
        return None

    page_html = (
        "<html><body style='margin:0;background:#fff'>"
        f"<div style='width:{width}px'>{svg}</div></body></html>"
    )
    try:
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch()
            try:
                page = browser.new_page(
                    viewport={"width": width, "height": height},
                    device_scale_factor=2,
                )
                page.route("**/*", lambda route: route.abort())
                page.set_content(page_html, wait_until="load")
                return page.locator("div").first.screenshot(type="png")
            finally:
                browser.close()
    except PlaywrightError:
        return None


def _write_block(document, block) -> None:
    if isinstance(block, Para):
        paragraph = document.add_paragraph()
        _write_runs(
            paragraph,
            block.runs,
            size=9.5 if block.note else 10.5,
            color=INK_SOFT if block.note else INK,
            italic=block.note,
        )
    elif isinstance(block, Sub):
        paragraph = document.add_paragraph(block.text, style="Heading 3")
        _keep_with_next(paragraph)
    elif isinstance(block, Bullets):
        for item in block.items:
            paragraph = document.add_paragraph(style="List Bullet")
            _write_runs(paragraph, item)
    elif isinstance(block, Table):
        _write_table(document, block)
    elif isinstance(block, Figure):
        _write_figure(document, block)
    elif isinstance(block, Callout):
        # Word has no boxed callout without a wrapper table, and a one-cell
        # table here would fight the page break. An indented rule-off block
        # reads the same and stays selectable text.
        for inner in block.blocks:
            _write_block(document, inner)
            if document.paragraphs:
                document.paragraphs[-1].paragraph_format.left_indent = Pt(14)
    else:  # pragma: no cover - the model is closed
        raise TypeError(f"unrenderable block: {block!r}")


# --- the document ------------------------------------------------------------


def _write_cover(document, doc: Document) -> None:
    kicker = document.add_paragraph()
    run = kicker.add_run(doc.kicker.upper())
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = INK_FAINT

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(doc.title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = INK

    system = document.add_paragraph()
    system.paragraph_format.left_indent = Pt(14)
    _left_border(system)
    run = system.add_run(doc.system_line)
    run.font.size = Pt(11)
    run.font.color.rgb = INK_SOFT

    meta = document.add_table(rows=0, cols=2)
    meta.autofit = False
    _left_rule(meta)
    for label, value in doc.meta:
        cells = meta.add_row().cells
        cells[0].width = Emu(int(CONTENT_W * 0.28))
        cells[1].width = Emu(int(CONTENT_W * 0.72))
        left = cells[0].paragraphs[0]
        run = left.add_run(label.upper())
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = INK_FAINT
        right = cells[1].paragraphs[0]
        run = right.add_run(value)
        run.font.size = Pt(10.5)


def _write_contents(document, doc: Document) -> None:
    document.add_paragraph("Table of Contents", style="Heading 1")
    note = document.add_paragraph()
    _write_runs(
        note,
        docshell.runs(
            "Word fills this in when the document opens — or right-click the table "
            "and choose Update Field."
        ),
        size=9,
        color=INK_FAINT,
        italic=True,
    )
    _field(
        document.add_paragraph(),
        ' TOC \\o "1-2" \\h \\z \\u ',
        "Right-click here and choose Update Field to build the contents.",
    )


def _write_control(document, doc: Document) -> None:
    document.add_paragraph("Document Control", style="Heading 1")
    if doc.change_note:
        _write_block(document, Para(docshell.runs(doc.change_note), note=True))
    _write_table(
        document,
        Table(
            ["Version", "Date", "Summary"],
            [
                [docshell.runs(row.version), docshell.runs(row.when), docshell.runs(row.summary)]
                for row in doc.control
            ],
            widths=(16, 22, 62),
        ),
    )


def render_docx(doc: Document) -> bytes:
    """The whole document as a Word file."""
    document = _WordDocument()
    _configure_styles(document)
    section = _configure_page(document)
    _running_chrome(section, doc.running_title)
    _update_fields_on_open(document)

    _write_cover(document, doc)

    # A page break, not a section break. A second section would need its own
    # header/footer references or rely on Word's inherit-when-absent rule, and
    # the running chrome silently vanishing after page 1 is exactly the kind of
    # defect nobody notices until the document is with the client.
    document.add_page_break()
    _write_contents(document, doc)

    document.add_page_break()
    _write_control(document, doc)

    for number, part, subs in docshell.numbered(doc.parts):
        document.add_paragraph(f"{number} {part.title}", style="Heading 1")
        for sub_number, part_section in subs:
            document.add_paragraph(f"{sub_number} {part_section.title}", style="Heading 2")
            for block in part_section.blocks:
                _write_block(document, block)

    if doc.closing:
        _write_block(document, Para(docshell.runs(doc.closing), note=True))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
