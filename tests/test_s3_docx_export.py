"""The Word export of the two hand-off documents.

These assert the machinery a reader never sees but Word depends on — heading
*styles* rather than heading-shaped text, the TOC field, the PAGE/NUMPAGES
fields, one section so the running chrome cannot vanish after page 1 — plus the
one behaviour that matters when the document is wrong: a run that produced no
source-control flow and no release notes must still come out numbered 1-2-3.
"""

from __future__ import annotations

import io
import zipfile
from datetime import date, datetime
from unittest.mock import patch

import pytest
from docx import Document as WordDocument

from s3_enhancement import docshell, docx_export
from s3_enhancement.designdoc import (
    build_design_document,
    build_release_record,
    render_document_docx,
    render_release_record_docx,
)
from s3_enhancement.docgen import ReleaseNoteSet
from s3_enhancement.release import DeploymentPlan, PlanStep, ReleaseRecord, SuiteEvidence

# A 1x1 PNG. Standing in for the rasterised change map keeps the suite off
# Chromium — the real conversion is exercised by hand, not on every run.
PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4"
    "890000000a49444154789c6300010000050001od0a2db40000000049454e44ae426082".replace(
        "od", "0d"
    )
)

SAMPLE_DOC = """### Internal Design Document for QA Handoff

#### 1. Summary
Prospects resolve to the Guest preference.

#### 2. Affected Areas
- **`repos/enroldirect/applicants.py`**: the policy value.

#### 3. Risk Areas
- **Backward compatibility**: members and guests must not move.

#### 4. Suggested QA Focus
- An eligible prospect, and the preference named on the decision.
"""


def _plan() -> DeploymentPlan:
    return DeploymentPlan(
        steps=[PlanStep(1, "Deploy", "Restart the service.", "run.sh")],
        rollback=[PlanStep(1, "Revert", "Restore the previous commit.", "")],
        order_reason="",
        service_order=["enroldirect"],
    )


def _record(**overrides) -> ReleaseRecord:
    base = dict(
        story_label="US-2026-045",
        ticket_key="AMS-1045",
        released_by="Priya Nair",
        generated_at=datetime(2026, 8, 4, 14, 32),
        changed_files=["repos/enroldirect/applicants.py"],
        criteria=[],
        matrix=None,
        evidence=[SuiteEvidence("Regression suite", True, 27, 27, "Human-authored.")],
        approvals=[],
        plan=_plan(),
        notes=None,
        diagram_svg="",
        diagram_caption="",
        unproven=[],
        branch=None,
    )
    base.update(overrides)
    return ReleaseRecord(**base)


def _design(**overrides) -> bytes:
    args = dict(
        story_label="US-2026-045",
        ticket_key="AMS-1045",
        source="repos/enroldirect",
        changed_files=["repos/enroldirect/applicants.py"],
        prepared_by="Ravi Kumar",
        today=date(2026, 8, 4),
    )
    args.update(overrides)
    return render_document_docx(SAMPLE_DOC, **args)


def _open(blob: bytes):
    return WordDocument(io.BytesIO(blob))


def _headings(document) -> list[tuple[str, str]]:
    return [
        (p.style.name, p.text)
        for p in document.paragraphs
        if p.style.name.startswith("Heading")
    ]


# --- the file itself ---------------------------------------------------------


def test_the_design_document_is_a_readable_word_package():
    blob = _design()
    package = zipfile.ZipFile(io.BytesIO(blob))
    assert package.testzip() is None
    assert "word/document.xml" in package.namelist()
    assert _open(blob).paragraphs[1].text == "Design Document — Ready for QA"


def test_parts_and_sections_use_real_heading_styles():
    """Not merely heading-shaped: a maroon 16pt run is invisible to Word's
    navigation pane and to the TOC field, which collects by style."""
    headings = _headings(_open(_design()))
    assert ("Heading 1", "1.0 Introduction") in headings
    assert ("Heading 2", "1.1 Purpose of this document") in headings
    assert ("Heading 1", "4.0 Sign-off") in headings


def test_the_contents_is_a_toc_field_word_is_asked_to_update():
    package = zipfile.ZipFile(io.BytesIO(_design()))
    body = package.read("word/document.xml").decode()
    assert 'TOC \\o "1-2"' in body
    assert "updateFields" in package.read("word/settings.xml").decode()


def test_the_footer_numbers_pages_with_fields_not_text():
    """Nothing in this repo can compute a page number, so the footer must ask
    Word for it rather than print one."""
    footer = zipfile.ZipFile(io.BytesIO(_design())).read("word/footer1.xml").decode()
    assert "PAGE" in footer and "NUMPAGES" in footer
    assert docshell.ORG in footer and docshell.CLASSIFICATION in footer


def test_the_document_has_one_section_so_the_chrome_cannot_stop_after_page_one():
    """A second section without its own header/footer references leaves the
    running chrome to Word's inherit-when-absent rule. One section, no rule."""
    body = zipfile.ZipFile(io.BytesIO(_design())).read("word/document.xml").decode()
    assert body.count("<w:sectPr") == 1
    assert body.count("headerReference") == 1


def test_the_running_header_names_the_document_and_its_ticket():
    header = zipfile.ZipFile(io.BytesIO(_design())).read("word/header1.xml").decode()
    assert "US-2026-045" in header and "AMS-1045" in header


# --- content ------------------------------------------------------------------


def test_the_cover_carries_the_document_identity():
    table = _open(_design()).tables[0]
    rows = {row.cells[0].text: row.cells[1].text for row in table.rows}
    assert rows["DOCUMENT ID"] == "AMS-1045-DD"
    assert rows["SOURCE"] == "repos/enroldirect"
    assert rows["CLASSIFICATION"] == docshell.CLASSIFICATION


def test_code_inside_bold_keeps_its_monospace_and_loses_its_backticks():
    """The model writes affected areas as ``**`path.py`**``. A non-recursive
    inline pass leaves the backticks in the text, which is invisible in HTML
    and glaring in Word."""
    document = _open(_design())
    bullet = next(p for p in document.paragraphs if "applicants.py" in p.text)
    assert "`" not in bullet.text
    mono = [r for r in bullet.runs if r.font.name == docx_export.MONO_FONT]
    assert mono and mono[0].bold


def test_the_change_map_is_embedded_when_it_can_be_rendered():
    with patch("s3_enhancement.docx_export._rasterise", return_value=PNG):
        blob = _design(diagram_svg='<svg viewBox="0 0 100 50"></svg>', diagram_caption="Map.")
    media = [n for n in zipfile.ZipFile(io.BytesIO(blob)).namelist() if "word/media/" in n]
    assert len(media) == 1


def test_a_missing_browser_says_so_rather_than_dropping_the_change_map():
    """Silently shipping a design document without its change map is worse than
    shipping one that says the diagram could not be rendered here."""
    with patch("s3_enhancement.docx_export._rasterise", return_value=None):
        blob = _design(diagram_svg='<svg viewBox="0 0 100 50"></svg>')
    text = "\n".join(p.text for p in _open(blob).paragraphs)
    assert "Change map not embedded" in text
    assert not [n for n in zipfile.ZipFile(io.BytesIO(blob)).namelist() if "word/media/" in n]


def test_model_text_cannot_break_out_of_the_document_xml():
    blob = render_document_docx(
        "**Summary**\n<w:p>alert</w:p> & <script>", story_label="s", ticket_key="AMS-1"
    )
    text = "\n".join(p.text for p in _open(blob).paragraphs)
    assert "<w:p>alert</w:p> & <script>" in text


# --- the release record --------------------------------------------------------


def test_the_release_record_numbers_its_parts_contiguously_when_sections_are_absent():
    """A run with no source-control flow and no release notes must not leave a
    hole where those parts would have been — which is why numbering happens at
    render time rather than being written into a skeleton."""
    document = _open(render_release_record_docx(_record(), source="repos/enroldirect"))
    parts = [text for style, text in _headings(document) if style == "Heading 1"]
    numbered = [p for p in parts if p[1:3] == ".0"]
    assert [p.split()[0] for p in numbered] == ["1.0", "2.0", "3.0", "4.0"]
    assert "Source control" not in "\n".join(p.text for p in document.paragraphs)


def test_the_release_record_adds_the_notes_part_when_a_model_produced_them():
    record = _record(notes=ReleaseNoteSet("Client.", "Ops.", "Guide."))
    document = _open(render_release_record_docx(record))
    parts = [t for s, t in _headings(document) if s == "Heading 1" and t[1:3] == ".0"]
    assert parts[-1] == "5.0 Release notes"


def test_a_failing_suite_is_coloured_as_a_failure_in_the_evidence_table():
    record = _record(
        evidence=[SuiteEvidence("Generated tests", False, 7, 9, "Two failures.")]
    )
    document = _open(render_release_record_docx(record))
    table = next(t for t in document.tables if t.rows[0].cells[0].text == "Suite")
    verdict = table.rows[1].cells[2]
    assert verdict.text == "FAIL"
    assert verdict.paragraphs[0].runs[0].font.color.rgb == docx_export.ACCENT


# --- the shared model ----------------------------------------------------------


def test_every_block_the_model_defines_can_be_written_to_word():
    """The vocabulary is closed on purpose. If a block type is added to
    `docshell` without teaching this renderer, it must fail here rather than in
    front of a client."""
    blocks = [
        docshell.Para(docshell.runs("p")),
        docshell.Sub("s"),
        docshell.Bullets([docshell.runs("b")]),
        docshell.Table(["h"], [[docshell.runs("c")]]),
        docshell.Callout([docshell.Para(docshell.runs("c"))]),
    ]
    document = docshell.Document(
        kicker="k",
        title="t",
        running_title="r",
        system_line="s",
        meta=[("A", "b")],
        control=[docshell.ControlRow("v1.0", "today", "why")],
        parts=[docshell.Part("Part", [docshell.Section("Section", blocks)])],
    )
    assert docx_export.render_docx(document).startswith(b"PK")


def test_an_unknown_block_is_refused_rather_than_skipped():
    document = docshell.Document(
        kicker="k",
        title="t",
        running_title="r",
        system_line="s",
        meta=[],
        control=[],
        parts=[docshell.Part("Part", [docshell.Section("Section", [object()])])],
    )
    with pytest.raises(TypeError, match="unrenderable block"):
        docx_export.render_docx(document)


def test_html_and_word_are_built_from_the_same_document():
    """The two renderers share a model so they cannot disagree; this pins that
    the builder is the single source, not that the outputs look alike."""
    built = build_design_document(SAMPLE_DOC, story_label="US-2026-045", ticket_key="AMS-1045")
    assert [p.title for p in built.parts] == [
        "Introduction",
        "The change",
        "Hand-off to QA",
        "Sign-off",
    ]
    assert build_release_record(_record()).title == "Release Record — Ready for Release"
